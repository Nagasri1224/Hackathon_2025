import os
from flask import Flask, request, jsonify, send_from_directory, render_template
import pandas as pd
from docx import Document
from werkzeug.utils import secure_filename

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'output'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)


def filter_publications(df, start_year, end_year):
    df = df[(df['Year'] >= start_year) & (df['Year'] <= end_year)]
    journals = df[df['Type'].str.lower() == 'journal']
    conferences = df[df['Type'].str.lower() == 'conference']
    return journals, conferences


def generate_word_summary(journals, conferences, output_path):
    doc = Document()
    doc.add_heading('Publication Summary Report', 0)

    def add_section(title, data):
        doc.add_heading(title, level=1)
        if data.empty:
            doc.add_paragraph("No records found.")
        else:
            for _, row in data.iterrows():
                doc.add_paragraph(f"{row['Year']} - {row['Faculty Name']}: \"{row['Title']}\" ({row['Venue']})")

    add_section('Journals', journals)
    add_section('Conferences', conferences)

    doc.save(output_path)


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/upload', methods=['POST'])
def upload_file():
    if 'excelFile' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files['excelFile']
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400

    try:
        filename = secure_filename(file.filename)
        basename = os.path.splitext(filename)[0]
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)

        df = pd.read_excel(filepath)

        start_year = request.form.get("startYear", type=int)
        end_year = request.form.get("endYear", type=int)

        journals, conferences = filter_publications(df, start_year, end_year)

        journal_filename = f"{basename}_journal_summary.xlsx"
        conference_filename = f"{basename}_conference_summary.xlsx"
        summary_filename = f"{basename}_summary.docx"

        journal_path = os.path.join(app.config['OUTPUT_FOLDER'], journal_filename)
        conference_path = os.path.join(app.config['OUTPUT_FOLDER'], conference_filename)
        summary_path = os.path.join(app.config['OUTPUT_FOLDER'], summary_filename)

        if not journals.empty:
            journals.to_excel(journal_path, index=False)
        if not conferences.empty:
            conferences.to_excel(conference_path, index=False)

        generate_word_summary(journals, conferences, summary_path)

        return jsonify({
            "journal": f"/output/{journal_filename}",
            "conference": f"/output/{conference_filename}",
            "summary": f"/output/{summary_filename}"
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/output/<filename>')
def download_file(filename):
    return send_from_directory(app.config['OUTPUT_FOLDER'], filename)


if __name__ == '__main__':
    app.run(debug=True)
